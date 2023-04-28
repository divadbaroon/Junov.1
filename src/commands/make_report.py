import docx
from ask_gpt import AskGPT

class CreateReport:
    
    def __init__(self):
        self.doc = docx.Document()
        self.ask_gpt = AskGPT()
    
    def ask_gpt_for_contents(self, prompt):
        
        contents = self.ask_gpt.ask_gpt(prompt)
        return contents
    
    def creeate_report(self, document_header, contents):

        # Add a heading
        self.doc.add_heading(document_header, level=1)

        # Add a paragraph with some text
        self.doc.add_paragraph(contents)

        # Save the document to a file
        self.doc.save('example.docx')
